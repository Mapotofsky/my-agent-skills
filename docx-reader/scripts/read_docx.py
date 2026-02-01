import argparse
import json
import os
import sys
from typing import Optional, Tuple
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def parse_bool(value: str) -> bool:
    value = value.lower()
    if value in ("true", "1", "yes", "y"):
        return True
    if value in ("false", "0", "no", "n"):
        return False
    raise ValueError("include-content 仅支持 true/false")


def normalize_range(total: int, start: Optional[int], end: Optional[int]) -> Tuple[int, int]:
    if total <= 0:
        return 0, 0
    if start is None:
        start = 1
    if end is None:
        end = total
    start = max(1, int(start))
    end = min(total, int(end))
    if start > end:
        return 0, 0
    return start - 1, end


def extract_docx(
    file_path: str,
    paragraph_start: Optional[int] = None,
    paragraph_end: Optional[int] = None,
    table_start: Optional[int] = None,
    table_end: Optional[int] = None,
    include_content: bool = True
) -> dict:
    if not os.path.isfile(file_path):
        return {
            "success": False,
            "file_path": file_path,
            "content": "",
            "statistics": {
                "paragraph_count": 0,
                "table_count": 0,
                "table_row_counts": [],
                "char_count": 0
            },
            "error": "文件不存在"
        }

    if not file_path.lower().endswith(".docx"):
        return {
            "success": False,
            "file_path": file_path,
            "content": "",
            "statistics": {
                "paragraph_count": 0,
                "table_count": 0,
                "table_row_counts": [],
                "char_count": 0
            },
            "error": "仅支持.docx格式"
        }

    try:
        document = Document(file_path)
    except PackageNotFoundError:
        return {
            "success": False,
            "file_path": file_path,
            "content": "",
            "statistics": {
                "paragraph_count": 0,
                "table_count": 0,
                "table_row_counts": [],
                "char_count": 0
            },
            "error": "无法读取docx文件"
        }
    except Exception as exc:
        return {
            "success": False,
            "file_path": file_path,
            "content": "",
            "statistics": {
                "paragraph_count": 0,
                "table_count": 0,
                "table_row_counts": [],
                "char_count": 0
            },
            "error": str(exc)
        }

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    tables = []
    table_row_counts = []

    for table in document.tables:
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_cells)
        tables.append(table_rows)
        table_row_counts.append(len(table_rows))

    paragraph_start_index, paragraph_end_index = normalize_range(
        len(paragraphs), paragraph_start, paragraph_end
    )
    table_start_index, table_end_index = normalize_range(len(tables), table_start, table_end)

    selected_paragraphs = paragraphs[paragraph_start_index:paragraph_end_index]
    selected_tables = tables[table_start_index:table_end_index]
    selected_table_row_counts = table_row_counts[table_start_index:table_end_index]

    table_text_blocks = []
    for table in selected_tables:
        rows = ["\t".join(cells) for cells in table]
        table_text_blocks.append("\n".join(rows))

    content_parts = []
    if selected_paragraphs:
        content_parts.append("\n".join(selected_paragraphs))
    if table_text_blocks:
        content_parts.append("\n\n".join(table_text_blocks))

    content = "\n\n".join(content_parts).strip()
    char_count = len(content)
    if not include_content:
        content = ""

    return {
        "success": True,
        "file_path": file_path,
        "content": content,
        "statistics": {
            "paragraph_count": len(selected_paragraphs),
            "table_count": len(selected_tables),
            "table_row_counts": selected_table_row_counts,
            "char_count": char_count
        },
        "error": None
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("file_path")
    parser.add_argument("--paragraph-start", type=int, default=None)
    parser.add_argument("--paragraph-end", type=int, default=None)
    parser.add_argument("--table-start", type=int, default=None)
    parser.add_argument("--table-end", type=int, default=None)
    parser.add_argument("--include-content", type=parse_bool, default=True)

    args = parser.parse_args()
    result = extract_docx(
        args.file_path,
        paragraph_start=args.paragraph_start,
        paragraph_end=args.paragraph_end,
        table_start=args.table_start,
        table_end=args.table_end,
        include_content=args.include_content
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))

    if not result["success"]:
        sys.exit(1)


if __name__ == "__main__":
    main()
